import streamlit as st
from typing import Optional
import PyPDF2
from docx import Document
import io
from PIL import Image


class FileProcessor:
    """Service for processing different file types and extracting text content"""

    def __init__(self):
        pass

    def extract_text(self, uploaded_file, file_type: str) -> Optional[str]:
        """
        Extract text content from uploaded file based on file type

        Args:
            uploaded_file: Streamlit uploaded file object
            file_type: Type of file (pdf, docx, etc.)

        Returns:
            Extracted text content or None if extraction fails
        """
        try:
            if file_type == 'pdf':
                return self._extract_from_pdf(uploaded_file)
            elif file_type == 'docx':
                return self._extract_from_docx(uploaded_file)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

        except Exception as e:
            st.error(f"Error extracting text from {uploaded_file.name}: {str(e)}")
            return None

    def _extract_from_pdf(self, uploaded_file) -> str:
        """Extract text from PDF file"""
        try:
            # Read the PDF file
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))

            text_content = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content += page.extract_text() + "\n"

            # Reset file pointer for potential reuse
            uploaded_file.seek(0)

            return text_content.strip()

        except Exception as e:
            # Try alternative PDF processing with pdfplumber if PyPDF2 fails
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
                    text_content = ""
                    for page in pdf.pages:
                        text_content += page.extract_text() + "\n"

                    uploaded_file.seek(0)
                    return text_content.strip()

            except ImportError:
                st.error("PDF processing failed. Please ensure the PDF contains selectable text.")
                raise e

    def _extract_from_docx(self, uploaded_file) -> str:
        """Extract text from DOCX file"""
        try:
            # Read the DOCX file
            doc = Document(io.BytesIO(uploaded_file.read()))

            text_content = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"

            # Reset file pointer for potential reuse
            uploaded_file.seek(0)

            return text_content.strip()

        except Exception as e:
            st.error(f"Error processing DOCX file: {str(e)}")
            raise e

    def validate_file_content(self, text_content: str) -> bool:
        """
        Validate that extracted text content is meaningful

        Args:
            text_content: Extracted text to validate

        Returns:
            True if content appears valid, False otherwise
        """
        if not text_content or not text_content.strip():
            return False

        # Check for minimum content length
        if len(text_content.strip()) < 10:
            return False

        # Check for reasonable character distribution
        # (to catch cases where extraction returned mostly symbols/artifacts)
        alphanumeric_chars = sum(1 for c in text_content if c.isalnum())
        total_chars = len(text_content)

        if total_chars > 0 and (alphanumeric_chars / total_chars) < 0.3:
            return False

        return True

    def get_file_info(self, uploaded_file) -> dict:
        """
        Get basic information about the uploaded file

        Args:
            uploaded_file: Streamlit uploaded file object

        Returns:
            Dictionary containing file information
        """
        return {
            'filename': uploaded_file.name,
            'size': uploaded_file.size,
            'type': uploaded_file.type,
            'file_extension': uploaded_file.name.split('.')[-1].lower()
        }
